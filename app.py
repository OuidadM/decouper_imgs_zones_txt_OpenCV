from flask import Flask, request, send_file, jsonify
import requests, base64, os, io
from markdown2 import markdown as md2html
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


app = Flask(__name__)

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")


def add_markdown_to_docx(markdown_text,temp_file, append=False):
    html = md2html(markdown_text)
    soup = BeautifulSoup(html, "html.parser")

    # Charger doc existant ou nouveau
    if append and os.path.exists(temp_file):
        doc = Document(temp_file)
        doc.add_page_break()
    else:
        doc = Document()

    for elem in soup.children:
        if elem.name == "h1":
            p = doc.add_paragraph(elem.get_text())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)

        elif elem.name == "p":
            text = elem.get_text()
            p = doc.add_paragraph(text)
            if text.startswith(">") or "[ALIGN=right]" in text:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif text.startswith("[ALIGN=center]"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif elem.name == "ul":
            for li in elem.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")

        elif elem.name == "ol":
            for li in elem.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")

    doc.save(temp_file)
    return temp_file

@app.route("/traduire", methods=["POST"])
def translate():
    if "image" not in request.files:
        return jsonify({"error": "Image manquante"}), 400
    image_bytes = request.files["image"].read()
    nom = request.args.get("nomFichier", "document")
    langue = "français" if nom.startswith("FR_") else "espagnol"
    bundle_index = int(request.args.get("bundle", "1"))
    temp_file = f"{nom}.docx"  # Fichier temporaire pour accumuler les pages

    # Encodage image
    image_data = base64.b64encode(image_bytes).decode("utf-8")

    # Prompt
    prompt_text = f"""
Voici une image d'un document administratif multilingue (français, arabe, anglais).
Traduis fidèlement tout le contenu visible en {langue}.
Respecte exactement la structure visuelle. Ne saute aucun élément.
"""

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": "https://test.local",
        "X-Title": "Traduction Document",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "anthropic/claude-3.5-sonnet",
        "max_tokens": 800,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
                ]
            }
        ]
    }

    response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
    data = response.json()

    if "choices" not in data:
        return jsonify({"error": "Erreur API", "details": data}), 500

    markdown_text = data["choices"][0]["message"]["content"]

    # Ajout dans le docx
    temp_file_path = add_markdown_to_docx(markdown_text,temp_file, append=(bundle_index > 1))

    # Si dernier bundle → renvoyer le fichier complet
    is_last = request.args.get("last", "false").lower() == "true"
    if is_last:
        return send_file(temp_file_path, as_attachment=True, download_name="traduction.docx")

    return jsonify({"status": f"Page {bundle_index} ajoutée avec succès."})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
